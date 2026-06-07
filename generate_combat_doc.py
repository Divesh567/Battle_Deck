from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name, size_pt, bold=False, color=None):
    s = styles[style_name]
    s.font.name      = font_name
    s.font.size      = Pt(size_pt)
    s.font.bold      = bold
    if color:
        s.font.color.rgb = RGBColor(*color)

set_style('Normal',    'Calibri', 11)
set_style('Heading 1', 'Calibri', 18, bold=True,  color=(0x1F, 0x49, 0x7D))
set_style('Heading 2', 'Calibri', 14, bold=True,  color=(0x2E, 0x74, 0xB5))
set_style('Heading 3', 'Calibri', 12, bold=True,  color=(0x20, 0x60, 0x9D))

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix + ' ')
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def note(text):
    """Italic grey note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(4)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title_p.add_run('Battle Deck')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
run2 = sub_p.add_run('Combat System Design')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta_p = doc.add_paragraph()
run3 = meta_p.add_run('For: Design Team   |   Status: In Development   |   Date: May 2026')
run3.font.size = Pt(10)
run3.italic = True
run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
h1('1. Overview')
body(
    'Battle Deck is a turn-based card game where the player builds a loadout before each fight, '
    'then uses a hand of cards drawn each turn to defeat an enemy. The core loop is:'
)
bullet('Select a Character, Class, and Weapon to define your deck and stats.')
bullet('Choose an enemy to fight.')
bullet('Play cards from your hand each turn, spending energy to do so.')
bullet('Outlast the opponent — reduce their HP to zero before yours hits zero.')
body(
    'The SimScene is the current sandbox environment where this loop can be tested '
    'end-to-end with any combination of assets.'
)

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 2. PRE-COMBAT: THE LOADOUT
# ═════════════════════════════════════════════════════════════════════════════
h1('2. Pre-Combat: The Loadout')
body('Before any fight begins the player makes three choices. Together these define both '
     'the player\'s stat block and the entire deck they will fight with.')

h2('Character (Entity)')
bullet('Determines the base stat block: max HP, base attack, magic power, starting energy, initiative, and max armor.')
bullet('Comes with a "starting deck" — a set of cards pre-assigned to that character.')

h2('Class / Skill Set')
bullet('Adds bonus stats on top of the character\'s base (HP, attack, magic power, initiative).')
bullet('Contributes "signature cards" — class-specific cards added to the deck.')
bullet('Defines weapon proficiencies: which weapon categories the class can use and at what mastery level.')

h2('Weapon')
bullet('Adds a flat attack bonus to the character\'s attack stat.')
bullet('Contributes "granted cards" — weapon-specific cards added to the deck.')
bullet('Requires the equipped class to have sufficient proficiency. If a class defines no proficiencies, all weapons are available.')

h2('Enemy')
bullet('The player also selects an opponent, each tagged with a tier: Minion, Elite, or Boss.')
bullet('Enemies have their own full stat block and a scripted set of moves called Intents.')

h2('The Resulting Deck')
body('The player\'s deck is assembled as: Starting Deck cards + Class Signature cards + Weapon-granted cards, then shuffled. '
     'Every loadout combination produces a distinct deck, which is the primary driver of build variety.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 3. INITIATIVE
# ═════════════════════════════════════════════════════════════════════════════
h1('3. Initiative — Who Goes First')
body('At the start of combat both combatants compare their Initiative stat. The higher value acts first. Ties go to the player.')
body('Initiative = Character base initiative + Class bonus initiative.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 4. THE TURN LOOP
# ═════════════════════════════════════════════════════════════════════════════
h1('4. The Turn Loop')

h2('Start of Turn (both sides)')
numbered('Armor resets to the combatant\'s max armor value.')
numbered('Energy resets to the combatant\'s starting energy value.')
numbered('Cards are drawn from the draw pile (player draws their configured hand size; enemies draw 3, though they don\'t play from a hand).')
numbered('Status effects tick — any status that deals damage at turn start (e.g. Bleeding) fires now. If the combatant dies from status damage, combat ends immediately.')
numbered('Stun check — if the combatant is Stunned, they lose their turn entirely and the turn passes to the opponent.')

h2('Player\'s Turn')
bullet('The player may play any number of cards from their hand, in any order, as long as they have enough energy to cover each card\'s cost.')
bullet('Cards that cannot be afforded are greyed out in the UI.')
bullet('The player ends their turn manually by pressing End Turn.')
bullet('Played cards move from the hand to the discard pile.')

h2('Enemy\'s Turn')
bullet('Enemies do not draw or play cards. Instead they execute a pre-authored "Intent" — a named action with one or more effects.')
bullet('By default enemies cycle through their intent list in fixed order. If the Shuffle Intents flag is on, they pick randomly instead.')
bullet('The enemy\'s next intent is always shown to the player in the UI before it fires — a core design pillar so the player can plan and react.')
bullet('After executing their intent, the enemy\'s turn ends automatically.')

h2('End of Turn (both sides)')
bullet('The active combatant discards their entire remaining hand.')
bullet('All status durations tick down by 1; any status that reaches 0 is removed.')
bullet('The turn passes to the opponent and the next turn begins.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 5. CARDS
# ═════════════════════════════════════════════════════════════════════════════
h1('5. Cards')
body('Every card has the following properties:')
bullet('Name and card text', bold_prefix='')
bullet('Energy cost — how much energy it takes to play. Greyed out if unaffordable.')
bullet('Card type — Attack, Skill, Power, or Reaction.')
bullet('Damage type — Physical, Fire, Ice, Lightning, Poison, Arcane, Holy, or Shadow. Non-damage cards have no type.')
bullet('Rarity — Common, Uncommon, Rare, Signature.')
bullet('Effects — an ordered list of one or more atomic effects that fire on play.')

h2('Deck Cycling')
body('When the draw pile is empty and a draw is needed, the entire discard pile is shuffled and becomes the new draw pile. '
     'Cards never leave the deck mid-fight under the current rules (a future Exhaust mechanic would change this).')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 6. EFFECTS SYSTEM
# ═════════════════════════════════════════════════════════════════════════════
h1('6. Effects System')
body('Each card is made up of one or more atomic "effects." An effect is an independent unit that does exactly one thing. '
     'Stacking multiple effects on a single card enables complex behaviour without special-casing cards in the engine.')

h2('Effect Types')
bullet('Deal Damage — to a single enemy, all enemies, or a random enemy.')
bullet('Heal — self or all allies.')
bullet('Gain Armor.')
bullet('Apply a Status — Bleeding, Burn, Poison, Freeze, Stun, Vulnerable, Weak, or Strength.')
bullet('Draw Cards.')
bullet('Discard Cards.')
bullet('Add Card to Hand.')
bullet('Exhaust Card — removes a card from the deck permanently for the fight.')
bullet('Gain Energy / Lose Energy.')
bullet('Conditional — checks a game state condition; fires one set of effects if true, another if false.')

h2('Target')
body('Each effect specifies who it applies to: Self, Single Enemy, All Enemies, Random Enemy, Single Ally, All Allies, or Everyone.')

h2('Base Value')
body('The raw number for the effect — damage dealt, HP healed, armor gained, cards drawn, etc.')

h2('Scaling Source (Designed — Not Yet Active)')
body('Effects can be configured to scale their value with a stat:')
bullet('Base Attack')
bullet('Magic Power')
bullet('Missing Health (stronger when closer to death)')
bullet('Cards in Hand')
bullet('Armor Amount')
bullet('Turns Elapsed')
note('This system is fully modelled in the data but the calculation is currently disabled in the engine.')

h2('Effect Flags')
bullet('Can Crit — this effect is eligible to trigger a critical hit.')
bullet('Pierce Armor — this effect bypasses the target\'s armor entirely.')
bullet('Trigger On-Hit Effects — this hit counts as a weapon hit for future on-hit procs (not yet implemented).')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 7. DAMAGE RESOLUTION
# ═════════════════════════════════════════════════════════════════════════════
h1('7. Damage Resolution')
body('When a damage effect fires, the engine calculates final damage in this exact order:')
numbered('Start with the base value of the effect.')
numbered('Weak modifier — if the attacker has the Weak status, multiply damage by 0.75 (–25%).')
numbered('Vulnerable modifier — if the target has the Vulnerable status, multiply damage by 1.5 (+50%).')
numbered('(Designed, inactive) Elemental Resistance — multiply by the target\'s resistance to that damage type.')
numbered('Critical Hit — if the effect can crit and a random roll beats the attacker\'s crit chance, multiply by 2.')
numbered('Round to the nearest integer. Clamp to a minimum of 1 — no hit can deal zero damage.')
numbered('Armor absorption — if the effect does not pierce armor and the target has armor, armor absorbs damage first. Armor is reduced by the amount absorbed; leftover damage hits HP.')
numbered('Subtract the final damage from the target\'s current HP.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 8. STATUS EFFECTS
# ═════════════════════════════════════════════════════════════════════════════
h1('8. Status Effects')
body('Statuses are tracked per combatant as a name paired with a stack or duration count.')

h2('Currently Implemented')
bullet('Bleeding (stacking damage) — at the start of the afflicted combatant\'s turn, they take damage equal to their current Bleeding stacks.')

h2('Designed — Not Yet Wired Up')
bullet('Weak — reduces all outgoing damage by 25%. The damage modifier exists in the engine; the card effect to apply it is commented out.')
bullet('Vulnerable — increases all incoming damage by 50%. Same situation.')
bullet('Stun — causes the combatant to lose their next turn. The turn-skip check is implemented; application is commented out.')
bullet('Poison — periodic damage similar to Bleeding, likely non-stacking.')
bullet('Burn — periodic fire-typed damage.')
bullet('Freeze — referenced in the condition system; exact behaviour not yet defined.')
bullet('Strength — flat bonus added to all outgoing damage/heals. The hook exists in the engine but is commented out.')

h2('Tick Behaviour')
body('At the end of every turn, all status values are decremented by 1. Any status that hits 0 is automatically removed. '
     'Statuses that deal damage (Bleeding) do so at the start of the turn before the combatant acts.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 9. CONDITIONAL EFFECTS
# ═════════════════════════════════════════════════════════════════════════════
h1('9. Conditional Effects')
body('A card effect can be of type Conditional, which checks a game state and branches into '
     '"then" effects if true or "else" effects if false. Both branches are themselves full effect arrays, '
     'enabling chains of conditions.')
body('Currently supported conditions:')
bullet('If target is Burning')
bullet('If target is Poisoned')
bullet('If target is Stunned')
bullet('If target is Frozen')
bullet('If target is Vulnerable')
bullet('If caster is at full HP')
bullet('If caster is at low HP (25% or below)')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 10. ENEMY DESIGN
# ═════════════════════════════════════════════════════════════════════════════
h1('10. Enemy Design')
body('Each enemy is a data asset containing:')
bullet('A full stat block: HP, attack, armor, initiative, crit chance, and elemental resistances.')
bullet('A tier: Minion, Elite, or Boss — used for UI labelling and future encounter design.')
bullet('An Intent Pattern — an ordered list of named moves. Each intent has a display name (shown to the player) and a list of effects that fire when it executes.')
bullet('Shuffle Intents flag — off means the enemy cycles their moves in a fixed, predictable order; on means they pick randomly.')
bullet('A loot table: a gold drop range and a pool of cards the player can be rewarded from.')
body('The enemy\'s next intent is always telegraphed to the player. This is a core design pillar: '
     'the player should never be surprised — they should always have information to react to.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 11. WHAT IS LEFT TO IMPLEMENT
# ═════════════════════════════════════════════════════════════════════════════
h1('11. What Is Left to Implement')
note('This section is for the development team but is included here so the design team understands '
     'which features are safe to design around vs. which are still pending.')

h2('Status Application')
body('Stun, Poison, Burn, Freeze, Vulnerable, Weak, and Strength are all defined in the effect type list '
     'but their application cases in the combat engine are commented out. '
     'Bleeding is the only status fully working end-to-end. The others need to be uncommented and validated.')

h2('Elemental Resistances')
body('Each entity already has resistance values for all eight damage types (Physical, Fire, Ice, etc.) '
     'and the lookup logic exists in the engine — but the resistance step in damage resolution is disabled. '
     'It needs to be turned on and the values tuned for balance.')

h2('Stat Scaling for Effects')
body('The entire ScalingSource system is modelled (Base Attack, Magic Power, Missing Health, '
     'Cards in Hand, Armor, Turns Elapsed) but the calculation is commented out. '
     'Currently all effects use only their flat base value.')

h2('Strength Status')
body('Strength is meant to add a flat damage bonus to all outgoing effects. '
     'The hook exists in the value calculation but is commented out.')

h2('Enemy Intent Preview Accuracy')
body('The UI always shows Intent #0 as the upcoming enemy action regardless of which intent '
     'the enemy is actually about to execute. The engine tracks this correctly internally — '
     'it just isn\'t exposed to the UI yet.')

h2('Card Manipulation Effects')
body('Discard Cards, Exhaust Card, and Add Card to Hand are defined as effect types '
     'but are not handled in the engine\'s switch statement. Playing a card with these effects does nothing.')

h2('Reaction Cards')
body('The Reaction card type exists, and data fields for reaction triggers are modelled (commented out). '
     'The logic to play a card outside the player\'s turn is not yet implemented.')

h2('Card Modifiers: Ethereal and Innate')
body('Ethereal (auto-discard at end of turn if unplayed) and Innate (always in the opening hand) '
     'are modelled on card data but commented out and not implemented.')

h2('Card Upgrades')
body('A reference to an "upgraded version" of each card exists on the card base class (commented out). '
     'No upgrade system is built yet.')

h2('Loadout Info Panel')
body('The loadout screen stat panels are partially commented out — they currently only show '
     'names, not the full stat breakdown that was designed (HP, ATK, proficiencies, etc.).')

h2('Multi-Target Logic')
body('All Enemies and Random Enemy targeting exist but currently both resolve to the single opponent. '
     'Once multi-enemy encounters are added, the target resolver needs to be updated.')

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 12. EXPANSION IDEAS
# ═════════════════════════════════════════════════════════════════════════════
h1('12. Expansion Ideas')

h2('Rogue-lite Run Structure')
body('The SimScene is a single-fight sandbox. The next step is a full run: '
     'a sequence of fights with escalating difficulty, optional rest and shop stops between encounters, '
     'and a persistent deck that the player grows by picking card rewards after each fight.')

h2('Card Rewards and Deck Building')
body('After each fight, offer three random cards to choose from — '
     'one could always be Rare or Signature. '
     'The enemy loot tables (card drop pools) are already built; they just need a reward-screen to surface them.')

h2('Card Upgrades')
body('Allow players to upgrade cards to a "+" version with improved values or added effects. '
     'The data hook is already modelled. A campfire or rest node in a run could offer one or two upgrades.')

h2('Relics / Passive Items')
body('Persistent passive items that silently modify how the deck behaves — '
     '"your first card each turn costs 0," "gain 1 armor whenever you draw a card," '
     '"Bleeding stacks deal double damage." These would be a list on the player\'s run state, '
     'checked at the relevant engine hooks.')

h2('Multi-Enemy Encounters')
body('Fights with two or three enemies where All Enemies and Random Enemy targeting become meaningful. '
     'Boss fights could open with a minion phase. '
     'This is the main prerequisite for making the multi-target effect types useful.')

h2('Status Effect Depth and Interactions')
body('Once the full status roster is live, design around interactions: '
     'a card that deals double damage to Frozen targets; '
     'a card that converts all Burn stacks into a burst of damage; '
     'Poison that stacks while Weakness makes it accelerate. '
     'The conditional system is already built to support this kind of synergy.')

h2('Scaling Builds')
body('Enabling the ScalingSource system unlocks build archetypes: '
     'a Warrior whose cards scale with Attack rewarding weapon investment; '
     'a Mage whose cards scale with Magic Power rewarding class synergy; '
     'a "Last Stand" card that scales with Missing Health for high-risk high-reward play. '
     'This adds a progression texture to the deck-building layer.')

h2('Enemy Phase Patterns')
body('Enemies that change behaviour mid-fight: below 50% HP they shuffle in a harder intent pattern '
     'or gain enrage stacks. Boss enemies could have a true second phase with different stats and moves.')

h2('Richer Conditional Conditions')
body('Expand the condition vocabulary: "If you have 5+ cards in hand," '
     '"If the enemy is below half HP," "If this is the first card you played this turn," '
     '"If you have played 3 Attack cards this turn." '
     'This turns the Conditional effect type into a puzzle designers can craft builds around.')

h2('Deck Archetypes and Synergy Tags')
body('Tag cards with keywords (Bleed, Frost, Momentum, Overcharge) and design cards that reference those tags. '
     'Example: Frost Burst deals +3 damage per Frost-tagged card played this turn. '
     'This gives the deck-building layer a rich synergy space beyond stat scaling.')

h2('Event Nodes')
body('Non-combat encounters in a run: '
     'a Merchant where you buy or remove cards, '
     'a Shrine where you gamble HP for a rare card, '
     'a Mystery Event with a narrative choice and a risk/reward outcome. '
     'None of these require changes to the combat engine — they live in the run-flow layer above it.')

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out_path = r"C:\Unity_Projects\Battle Deck\Battle-Deck\Battle Deck — Combat System Design.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
